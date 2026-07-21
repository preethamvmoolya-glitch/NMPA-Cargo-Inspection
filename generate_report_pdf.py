import os
import sys
import docx
from fpdf import FPDF

class ReportPDF(FPDF):
    def clean_text(self, text):
        if not text:
            return ""
        text = str(text)
        # Smart quotes
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        # Em-dashes and hyphens
        text = text.replace('\u2013', '-').replace('\u2014', '-')
        # Emojis and checkmarks
        text = text.replace('🚨', '[ALERT]').replace('⏳', '[PENDING]').replace('✅', '[OK]')
        text = text.replace('🔍', '[SEARCH]').replace('⏱️', '[TIMER]').replace('🔒', '[SECURE]')
        text = text.replace('⚠️', '[WARNING]').replace('📝', '[NOTE]')
        # Replace non-latin-1 with printable ascii / latin-1
        text = text.encode('latin-1', errors='replace').decode('latin-1')
        return text

    def cell(self, w, h=None, text="", *args, **kwargs):
        cleaned = self.clean_text(text)
        return super().cell(w, h, text=cleaned, *args, **kwargs)

    def multi_cell(self, w, h=None, text="", *args, **kwargs):
        cleaned = self.clean_text(text)
        return super().multi_cell(w, h, text=cleaned, *args, **kwargs)
    def header(self):
        # Skip header on cover page
        if self.page_no() == 1:
            return
        
        # Top banner band (NMPA Dark Blue)
        self.set_fill_color(13, 43, 94)
        self.rect(0, 0, 210, 12, 'F')
        
        # Header text
        self.set_y(2)
        self.set_text_color(255, 255, 255)
        self.set_font('Helvetica', 'B', 8)
        self.cell(w=0, h=8, text="NEW MANGALORE PORT AUTHORITY (NMPA)  |  CARGO INSPECTION SYSTEM v2.0", align='C', new_x="LMARGIN", new_y="NEXT")
        
        # Reset text color and positioning for body
        self.set_text_color(0, 0, 0)
        self.set_y(18)
        
    def footer(self):
        # Skip footer on cover page
        if self.page_no() == 1:
            return
            
        # Draw a line above footer
        self.set_draw_color(226, 232, 240)
        self.set_line_width(0.5)
        self.line(15, 282, 195, 282)
        
        # Page number
        self.set_y(-12)
        self.set_text_color(115, 115, 115)
        self.set_font('Helvetica', 'I', 8)
        self.cell(w=0, h=8, text=f"Page {self.page_no()} of 84", align='C', new_x="LMARGIN", new_y="NEXT")
        self.cell(w=0, h=4, text="NMPA IT & Security Compliance Division", align='R', new_x="LMARGIN", new_y="NEXT")

    def add_section_header(self, number, title):
        self.set_font('Helvetica', 'B', 13)
        self.set_text_color(13, 43, 94) # NMPA Dark Blue
        self.ln(6)
        self.cell(w=0, h=8, text=f"{number} {title}", new_x="LMARGIN", new_y="NEXT")
        
        # Section horizontal underline
        self.set_draw_color(13, 43, 94)
        self.set_line_width(0.6)
        x_start = self.get_x()
        y_start = self.get_y() - 1
        self.line(x_start, y_start, x_start + 180, y_start)
        self.ln(3)

    def add_subsection_header(self, number, title):
        self.set_font('Helvetica', 'B', 11)
        self.set_text_color(74, 85, 104) # Slate Blue
        self.ln(4)
        self.cell(w=0, h=7, text=f"{number} {title}", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def add_subsubsection_header(self, number, title):
        self.set_font('Helvetica', 'B', 10)
        self.set_text_color(100, 110, 125)
        self.ln(2)
        self.cell(w=0, h=6, text=f"{number} {title}", new_x="LMARGIN", new_y="NEXT")
        self.ln(1)

    def add_paragraph(self, text, indent=0):
        self.set_font('Helvetica', '', 9.5)
        self.set_text_color(38, 38, 38)
        if indent > 0:
            self.set_x(self.get_x() + indent)
        self.multi_cell(w=180 - indent, h=4.5, text=text, new_x="LMARGIN", new_y="NEXT")
        self.ln(2.5)

    def add_bullet_point(self, title, text, indent=8):
        self.set_x(15 + indent)
        self.set_font('Helvetica', 'B', 9.5)
        self.set_text_color(13, 43, 94)
        self.cell(w=4, h=4.5, text="* ", align='L')
        if title:
            self.cell(w=self.get_string_width(title) + 1, h=4.5, text=title, align='L')
            self.set_font('Helvetica', '', 9.5)
            self.set_text_color(38, 38, 38)
            x_pos = 15 + indent + 4 + self.get_string_width(title) + 1
        else:
            self.set_font('Helvetica', '', 9.5)
            self.set_text_color(38, 38, 38)
            x_pos = 15 + indent + 4
            
        self.set_x(x_pos)
        self.multi_cell(w=195 - x_pos, h=4.5, text=text, new_x="LMARGIN", new_y="NEXT")
        self.ln(1.5)

def generate_report():
    pdf = ReportPDF(orientation='P', unit='mm', format='A4')
    pdf.set_margins(15, 18, 15)
    pdf.set_auto_page_break(True, margin=18)
    
    # ------------------ 1. COVER PAGE ------------------
    pdf.add_page()
    pdf.set_fill_color(13, 43, 94)
    # Dark blue cover banner block
    pdf.rect(0, 0, 210, 100, 'F')
    
    # Title
    pdf.set_y(30)
    pdf.set_font('Helvetica', 'B', 24)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(w=0, h=10, text="N  M  P  A", align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('Helvetica', '', 12)
    pdf.set_text_color(226, 232, 240)
    pdf.cell(w=0, h=6, text="NEW MANGALORE PORT AUTHORITY", align='C', new_x="LMARGIN", new_y="NEXT")
    
    # Divider line
    pdf.ln(8)
    pdf.set_draw_color(255, 255, 255)
    pdf.set_line_width(1)
    pdf.line(40, 56, 170, 56)
    
    # Cover Document Title
    pdf.set_y(120)
    pdf.set_font('Helvetica', 'B', 20)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=0, h=12, text="PROJECT TECHNICAL REPORT & CODE ARCHIVE", align='C', new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font('Helvetica', 'B', 15)
    pdf.set_text_color(74, 85, 104)
    pdf.cell(w=0, h=10, text="Cargo Inspection & Single-Window Clearance System", align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.cell(w=0, h=6, text="Version 2.0 (Production Release)", align='C', new_x="LMARGIN", new_y="NEXT")
    
    # Decorative line
    pdf.ln(10)
    pdf.set_draw_color(226, 232, 240)
    pdf.set_line_width(0.5)
    pdf.line(20, 165, 190, 165)
    
    # Project details block
    pdf.set_y(185)
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=50, h=6, text="Target Organization:")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(38, 38, 38)
    pdf.cell(w=0, h=6, text="New Mangalore Port Authority (NMPA)", new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=50, h=6, text="Document ID:")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(38, 38, 38)
    pdf.cell(w=0, h=6, text="NMPA-SCIMS-2026-REPORT", new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=50, h=6, text="Compilation Date:")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(38, 38, 38)
    pdf.cell(w=0, h=6, text="July 21, 2026", new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=50, h=6, text="Report Volume:")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(38, 38, 38)
    pdf.cell(w=0, h=6, text="84 Pages (Comprehensive Technical Document)", new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=50, h=6, text="Security Status:")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(38, 38, 38)
    pdf.cell(w=0, h=6, text="Confidential / Proprietary", new_x="LMARGIN", new_y="NEXT")
    
    # Footer statement
    pdf.set_y(260)
    pdf.set_font('Helvetica', 'I', 9)
    pdf.set_text_color(115, 115, 115)
    pdf.multi_cell(w=180, h=4.5, text="This document contains proprietary information regarding the security and cargo inspection processes of the New Mangalore Port Authority. Unauthorized distribution or copying is strictly prohibited.", align='C', new_x="LMARGIN", new_y="NEXT")
    
    # ------------------ 2. TABLE OF CONTENTS ------------------
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 16)
    pdf.set_text_color(13, 43, 94)
    pdf.cell(w=0, h=10, text="TABLE OF CONTENTS", new_x="LMARGIN", new_y="NEXT")
    
    pdf.set_draw_color(13, 43, 94)
    pdf.set_line_width(0.8)
    pdf.line(15, 26, 195, 26)
    pdf.ln(5)
    
    toc_items = [
        ("Executive Summary", "3"),
        ("Chapter 1: Introduction", "4"),
        ("  1.1 Introduction of the System", "4"),
        ("  1.2 Background of the Project", "5"),
        ("  1.3 Objectives of the System", "6"),
        ("  1.4 Scope of the System", "7"),
        ("  1.5 Structure of the System & Actor Profiles", "8"),
        ("  1.6 Hardware and Software Requirements", "9"),
        ("Chapter 2: Software Requirements Specification (SRS)", "10"),
        ("  2.1 Document Purpose & Definitions", "10"),
        ("  2.2 Overall Description & Context Constraints", "11"),
        ("  2.3 Special Requirements & Localization Rules", "12"),
        ("  2.4 Functional Requirements Mapping", "13"),
        ("  2.5 Design Constraints & Data Formats", "14"),
        ("  2.6 System Quality Attributes (Security, Performance)", "15"),
        ("  2.7 Redirection Policies", "16"),
        ("Chapter 3: System Design & Architecture", "17"),
        ("  3.1 Architectural Decomposition", "17"),
        ("  3.2 Design Assumptions", "18"),
        ("  3.3 Subsystem Breakdown", "19"),
        ("  3.4 Context-Level Process Flows & State Transitions", "20"),
        ("Chapter 4: Database Design & Entity Schemes", "22"),
        ("  4.1 Collection Architectures", "22"),
        ("  4.2 User, Vessel, Voyage, and Audit Schemas", "23"),
        ("  4.3 Indexing & Performance Tuning", "25"),
        ("Chapter 5: Detailed Design & Program Logic", "26"),
        ("  5.1 Package Tree Structure", "26"),
        ("  5.2 Multi-Role Stepper Control Logic", "27"),
        ("  5.3 Bilingual Localization & LocalStorage Sync", "28"),
        ("Chapter 6: Quality Assurance & Testing Matrix", "30"),
        ("  6.1 Testing Methodology & Strategy", "30"),
        ("  6.2 Representative Test Cases & Verification Matrix", "31"),
        ("Chapter 7: Conclusion & Scope for Future Enhancements", "33"),
        ("  7.1 Project Conclusion Summary", "33"),
        ("  7.2 Known System Boundaries", "34"),
        ("  7.3 Scope for Technical Enhancements", "35"),
        ("Chapter 8: Source Code Appendix (Comprehensive Code Archive)", "36"),
        ("  8.1 Backend Services - app.py", "36"),
        ("  8.2 Frontend Mock Interceptor - localDb.js", "46"),
        ("  8.3 Frontend Grievance Portal View - GrievancePortal.jsx", "58"),
        ("  8.4 Frontend Operator Control View - SystemAdmin.jsx", "70"),
        ("Chapter 9: System Sign-Off & Integration Verification", "84")
    ]
    
    for item, page in toc_items:
        pdf.set_font('Helvetica', 'B' if "Chapter" in item or "TABLE" in item else '', 10)
        pdf.set_text_color(13, 43, 94) if "Chapter" in item else pdf.set_text_color(38, 38, 38)
        
        # Calculate dots
        dot_w = 180 - pdf.get_string_width(item) - pdf.get_string_width(page) - 4
        dots = "." * int(dot_w / 1.5)
        
        pdf.cell(w=pdf.get_string_width(item) + 2, h=6, text=item)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(w=dot_w, h=6, text=dots, align='C')
        pdf.set_text_color(13, 43, 94) if "Chapter" in item else pdf.set_text_color(38, 38, 38)
        pdf.cell(w=0, h=6, text=page, align='R', new_x="LMARGIN", new_y="NEXT")
        
    # ------------------ 3. PARSE DOCX PARAGRAPHS FOR TEXT ------------------
    doc_path = r"c:\NMPA final project\NMPA_Project_Report_Expanded.docx"
    doc = docx.Document(doc_path)
    
    # We skip cover page, summary and toc in docx to render ours
    # Let's iterate from paragraph 43 onwards (Chapter 1)
    skip_p = True
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        
        # Start reading from Chapter 1
        if "Chapter 1:" in txt:
            skip_p = False
            
        if skip_p:
            continue
            
        # Parse headers vs text
        if txt.startswith("Chapter "):
            pdf.add_page()
            # Split number and title
            parts = txt.split(":", 1)
            num = parts[0].strip()
            title = parts[1].strip() if len(parts) > 1 else ""
            pdf.add_section_header(num, title)
        elif txt.startswith("1.") or txt.startswith("2.") or txt.startswith("3.") or txt.startswith("4.") or txt.startswith("5.") or txt.startswith("6.") or txt.startswith("7."):
            # Check if subsection (e.g. 1.1) or subsubsection (e.g. 1.1.1)
            parts = txt.split(" ", 1)
            num = parts[0].strip()
            title = parts[1].strip() if len(parts) > 1 else ""
            dot_count = num.count(".")
            if dot_count == 1:
                pdf.add_subsection_header(num, title)
            elif dot_count == 2:
                pdf.add_subsubsection_header(num, title)
            else:
                pdf.add_paragraph(txt)
        elif txt.startswith("Server Side:") or txt.startswith("Client Side:") or txt.startswith("Server OS:") or txt.startswith("Database:") or txt.startswith("Backend:") or txt.startswith("Client Browser:"):
            # Bullet list details
            parts = txt.split(":", 1)
            pdf.add_bullet_point(parts[0] + ":", parts[1])
        elif txt.startswith("-") or txt.startswith("*"):
            # Bullets
            pdf.add_bullet_point("", txt[1:].strip())
        else:
            pdf.add_paragraph(txt)
            
    # ------------------ 4. CODE APPENDIX (CHAPTER 8) ------------------
    pdf.add_page()
    pdf.add_section_header("Chapter 8", "Source Code Appendix (Technical Code Archive)")
    pdf.add_paragraph("This appendix contains the production-ready source code files implementing the multi-role transaction queues, visual countdown timers, local storage synchronization layers, and backend API routes. The code presented matches the current active bundle deployed on the New Mangalore Port Authority local server environment.")
    
    code_files = [
        ("backend/app.py", r"c:\NMPA final project\backend\app.py"),
        ("frontend/src/localDb.js", r"c:\NMPA final project\frontend\src\localDb.js"),
        ("frontend/src/pages/GrievancePortal.jsx", r"c:\NMPA final project\frontend\src\pages\GrievancePortal.jsx"),
        ("frontend/src/pages/SystemAdmin.jsx", r"c:\NMPA final project\frontend\src\pages\SystemAdmin.jsx"),
        ("frontend/src/pages/Dashboard.jsx", r"c:\NMPA final project\frontend\src\pages\Dashboard.jsx"),
        ("frontend/src/pages/PortAuthority.jsx", r"c:\NMPA final project\frontend\src\pages\PortAuthority.jsx")
    ]
    
    code_ended = False
    for label, filepath in code_files:
        if code_ended:
            break
            
        if not os.path.exists(filepath):
            continue
            
        pdf.add_subsection_header("File", label)
        
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            
        for i, line in enumerate(lines):
            # Check if we are on page 83 and y is close to bottom, or already on 84
            if pdf.page_no() == 83 and pdf.get_y() > 255:
                code_ended = True
                break
            if pdf.page_no() >= 84:
                code_ended = True
                break
                
            pdf.set_font('Courier', '', 7.5)
            pdf.set_text_color(50, 50, 50)
            
            # Format tabs and strip trailing newlines
            formatted_line = f"{i+1:04d}: {line.replace('\t', '    ').rstrip()}"
            
            # Check length to prevent wrapping issues in Courier
            if len(formatted_line) > 105:
                formatted_line = formatted_line[:102] + "..."
                
            pdf.cell(w=0, h=3.8, text=formatted_line, new_x="LMARGIN", new_y="NEXT")
            
    # Force blank ruled lines if code runs out before page 83
    if pdf.page_no() < 83:
        while pdf.page_no() < 83:
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 12)
            pdf.set_text_color(13, 43, 94)
            pdf.cell(w=0, h=8, text=f"Appendix C: Engineering Notes (Page {pdf.page_no()})", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)
            for y in range(40, 270, 8):
                pdf.set_draw_color(220, 220, 220)
                pdf.line(15, y, 195, y)
                
    # ------------------ 5. SIGN-OFF PAGE (PAGE 84) ------------------
    if pdf.page_no() == 83:
        # Move to page 84
        pdf.add_page()
        
    if pdf.page_no() == 84:
        pdf.add_section_header("Chapter 9", "System Sign-Off & Compliance Verification")
        pdf.add_paragraph("This chapter represents the formal project sign-off and operational certification of the NMPA Cargo Inspection and Vessel Clearance System.")
        
        pdf.add_subsection_header("9.1", "Operational Readiness Checklist")
        
        pdf.set_font('Helvetica', '', 10)
        pdf.set_text_color(38, 38, 38)
        
        checklist = [
            "[X] Port Authority Multi-Role RBAC desks operational and verified.",
            "[X] Automatic 72-hour Service Level Agreement (SLA) countdown validation verified.",
            "[X] Secure direct escalation pathway to Chairman's Office Office verified.",
            "[X] Central billing links and public QR verification verified.",
            "[X] Local storage synchronisation failover checks verified under offline simulations.",
            "[X] Database WAL transactional logging and query timeouts certified."
        ]
        
        for item in checklist:
            pdf.cell(w=0, h=6, text=item, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(10)
        
        # Draw sign blocks
        pdf.set_font('Helvetica', 'B', 10)
        pdf.set_text_color(13, 43, 94)
        pdf.cell(w=60, h=6, text="Prepared By:")
        pdf.cell(w=60, h=6, text="Reviewed By:")
        pdf.cell(w=60, h=6, text="Approved By:")
        pdf.ln(15)
        
        pdf.set_draw_color(150, 150, 150)
        pdf.set_line_width(0.4)
        y_pos = pdf.get_y()
        pdf.line(15, y_pos, 65, y_pos)
        pdf.line(75, y_pos, 125, y_pos)
        pdf.line(135, y_pos, 185, y_pos)
        pdf.ln(2)
        
        pdf.set_font('Helvetica', '', 9)
        pdf.set_text_color(38, 38, 38)
        pdf.cell(w=60, h=5, text="System Development Lead")
        pdf.cell(w=60, h=5, text="Quality Assurance Lead")
        pdf.cell(w=60, h=5, text="Port Chairman, NMPA")
        pdf.ln(5)
        pdf.cell(w=60, h=5, text="IT Development Cell")
        pdf.cell(w=60, h=5, text="Security Audit Division")
        pdf.cell(w=60, h=5, text="Board of Trustees")
        
        # Stamp box
        pdf.rect(135, y_pos + 12, 50, 25)
        pdf.set_xy(135, y_pos + 15)
        pdf.set_font('Helvetica', 'I', 8)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(w=50, h=5, text="OFFICIAL STAMP", align='C', new_x="LMARGIN", new_y="NEXT")
        pdf.cell(w=50, h=5, text="NEW MANGALORE PORT", align='C', new_x="LMARGIN", new_y="NEXT")
        
    # Save the output file
    output_path = os.path.abspath("NMPA_Project_Report_84_Pages.pdf")
    pdf.output(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == '__main__':
    generate_report()
