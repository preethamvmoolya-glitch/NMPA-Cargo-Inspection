import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def clean_text(text):
    if not text:
        return ""
    text = str(text)
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2018', "'").replace('\u2019', "'")
    text = text.replace('\u2013', '-').replace('\u2014', '-')
    text = text.replace('🚨', '[ALERT]').replace('⏳', '[PENDING]').replace('✅', '[OK]')
    text = text.replace('🔍', '[SEARCH]').replace('⏱️', '[TIMER]').replace('🔒', '[SECURE]')
    text = text.replace('⚠️', '[WARNING]').replace('📝', '[NOTE]')
    # Replace other problematic unicode characters with ASCII replacements
    text = text.encode('ascii', errors='replace').decode('ascii')
    return text

def add_styled_heading(doc, text, level):
    clean_txt = clean_text(text)
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(clean_txt)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(13, 43, 94) # NMPA Dark Blue
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(clean_txt)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(74, 85, 104) # Slate Gray/Blue
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(clean_txt)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(100, 110, 125)
    return p

def add_styled_paragraph(doc, text):
    clean_txt = clean_text(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(clean_txt)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(38, 38, 38)
    return p

def add_styled_bullet(doc, title, text):
    clean_title = clean_text(title)
    clean_txt = clean_text(text)
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if clean_title:
        r_title = p.add_run(clean_title + " ")
        r_title.bold = True
        r_title.font.name = 'Times New Roman'
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(13, 43, 94)
        
    r_text = p.add_run(clean_txt)
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(38, 38, 38)
    return p

def add_code_block(doc, filename, code_lines):
    # Add title paragraph
    p_title = doc.add_paragraph()
    p_title.paragraph_format.keep_with_next = True
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(f"Source Code: {filename}")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(74, 85, 104)

    # Chunk lines to speed up paragraph generation (150 lines per chunk)
    chunk_size = 150
    for chunk_start in range(0, len(code_lines), chunk_size):
        chunk_lines = code_lines[chunk_start:chunk_start+chunk_size]
        formatted_chunk = []
        for idx, line in enumerate(chunk_lines):
            line_num = chunk_start + idx + 1
            formatted_line = f"{line_num:04d}: {line.replace('\t', '    ').rstrip()}"
            if len(formatted_line) > 105:
                formatted_line = formatted_line[:102] + "..."
            formatted_chunk.append(formatted_line)
            
        p_code = doc.add_paragraph()
        p_code.paragraph_format.space_before = Pt(0)
        p_code.paragraph_format.space_after = Pt(2)
        p_code.paragraph_format.line_spacing = 1.0
        
        run_code = p_code.add_run("\n".join(formatted_chunk))
        run_code.font.name = 'Courier New'
        run_code.font.size = Pt(8)
        run_code.font.color.rgb = RGBColor(70, 70, 70)

def add_ui_screenshot(doc, img_path, caption_text):
    if not os.path.exists(img_path):
        # Fallback block description if image is missing
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[UI SCREENSHOT ARCHIVE: {os.path.basename(img_path)}]")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(150, 0, 0)
        return
        
    # Insert screenshot
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(10)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.paragraph_format.keep_with_next = True
    
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(5.8))
    
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    
    run_cap = p_cap.add_run(clean_text(caption_text))
    run_cap.italic = True
    run_cap.font.name = 'Times New Roman'
    run_cap.font.size = Pt(9.5)
    run_cap.font.color.rgb = RGBColor(100, 100, 100)

def generate_report_docx():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # ------------------ 1. COVER PAGE ------------------
    # Large Space
    for _ in range(3):
        doc.add_paragraph()
        
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_org = p_org.add_run("NEW MANGALORE PORT AUTHORITY\nIT & Security Compliance Division")
    r_org.font.name = 'Arial'
    r_org.font.size = Pt(12)
    r_org.bold = True
    r_org.font.color.rgb = RGBColor(13, 43, 94)
    
    for _ in range(2):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("PROJECT TECHNICAL REPORT & COMPREHENSIVE CODE ARCHIVE")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(13, 43, 94)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Cargo Inspection & Single-Window Clearance Portal\nVersion 2.0 (Production Release Bundle)")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(14)
    r_sub.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Meta Block
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run(
        "ORGANIZATION: NEW MANGALORE PORT AUTHORITY\n"
        "SECTOR: MINISTRY OF PORTS, SHIPPING AND WATERWAYS\n"
        "DOCUMENT CLASS: UNIVERSITY-LEVEL PROJECT DOCUMENTATION\n"
        "DATE OF COMPILATION: JULY 21, 2026\n"
        "REPORT VOLUME: 84+ PAGES (STANDARD COMPLIANCE STRUCTURE)\n"
        "DEVELOPER CELL: NMPA IT DEVELOPMENT & OPERATIONS CENTRE"
    )
    r_meta.font.name = 'Times New Roman'
    r_meta.font.size = Pt(10.5)
    r_meta.font.color.rgb = RGBColor(50, 50, 50)
    
    doc.add_page_break()
    
    # ------------------ 2. TABLE OF CONTENTS ------------------
    p_toc_title = doc.add_paragraph()
    r_toc_title = p_toc_title.add_run("TABLE OF CONTENTS")
    r_toc_title.font.name = 'Arial'
    r_toc_title.font.size = Pt(16)
    r_toc_title.bold = True
    r_toc_title.font.color.rgb = RGBColor(13, 43, 94)
    p_toc_title.paragraph_format.space_after = Pt(12)
    
    toc_items = [
        ("Executive Summary", "3"),
        ("Chapter 1: Introduction", "4"),
        ("  1.1 Introduction of the System", "4"),
        ("  1.2 Background of the Project", "5"),
        ("  1.3 Objectives of the System", "6"),
        ("  1.4 Scope of the System", "7"),
        ("  1.5 Structure of the System & Actor Profiles", "8"),
        ("  1.6 Hardware and Software Requirements", "9"),
        ("Chapter 2: Software Requirements Specification (SRS)", "10"),
        ("  2.1 Document Purpose & Definitions", "10"),
        ("  2.2 Overall Description & Context Constraints", "11"),
        ("  2.3 Special Requirements & Localization Rules", "12"),
        ("  2.4 Functional Requirements Mapping", "13"),
        ("  2.5 Design Constraints & Data Formats", "14"),
        ("  2.6 System Quality Attributes (Security, Performance)", "15"),
        ("  2.7 Redirection Policies", "16"),
        ("Chapter 3: System Design & Architecture", "17"),
        ("  3.1 Architectural Decomposition", "17"),
        ("  3.2 Design Assumptions", "18"),
        ("  3.3 Subsystem Breakdown", "19"),
        ("  3.4 Context-Level Process Flows & State Transitions", "20"),
        ("Chapter 4: Database Design & Entity Schemes", "22"),
        ("  4.1 Collection Architectures", "22"),
        ("  4.2 User, Vessel, Voyage, and Audit Schemas", "23"),
        ("  4.3 Indexing & Performance Tuning", "25"),
        ("Chapter 5: Detailed Design & Program Logic", "26"),
        ("  5.1 Package Tree Structure", "26"),
        ("  5.2 Multi-Role Stepper Control Logic", "27"),
        ("  5.3 Bilingual Localization & LocalStorage Sync", "28"),
        ("  5.4 User Interface Gallery & Screenshot References", "30"),
        ("Chapter 6: Quality Assurance & Testing Matrix", "33"),
        ("  6.1 Testing Methodology & Strategy", "33"),
        ("  6.2 Representative Test Cases & Verification Matrix", "34"),
        ("Chapter 7: Conclusion & Scope for Future Enhancements", "36"),
        ("  7.1 Project Conclusion Summary", "36"),
        ("  7.2 Known System Boundaries", "37"),
        ("  7.3 Scope for Technical Enhancements", "38"),
        ("Chapter 8: Source Code Appendix (Comprehensive Code Archive)", "39"),
        ("  8.1 Backend Services - app.py", "39"),
        ("  8.2 Frontend Mock Interceptor - localDb.js", "48"),
        ("  8.3 Frontend Grievance Portal View - GrievancePortal.jsx", "59"),
        ("  8.4 Frontend Operator Control View - SystemAdmin.jsx", "71"),
        ("Chapter 9: System Sign-Off & Integration Verification", "84")
    ]
    
    for item, page in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_after = Pt(3)
        p_toc.paragraph_format.line_spacing = 1.0
        
        r_item = p_toc.add_run(item)
        r_item.font.name = 'Arial' if "Chapter" in item else 'Times New Roman'
        r_item.font.size = Pt(10)
        r_item.bold = "Chapter" in item
        r_item.font.color.rgb = RGBColor(13, 43, 94) if "Chapter" in item else RGBColor(38, 38, 38)
        
        # Dots
        dot_count = 100 - len(item) - len(page)
        dots = " " + ("." * max(5, dot_count)) + " "
        r_dots = p_toc.add_run(dots)
        r_dots.font.name = 'Courier New'
        r_dots.font.size = Pt(9.5)
        r_dots.font.color.rgb = RGBColor(150, 150, 150)
        
        r_page = p_toc.add_run(page)
        r_page.font.name = 'Times New Roman'
        r_page.font.size = Pt(10)
        r_page.bold = "Chapter" in item
        r_page.font.color.rgb = RGBColor(13, 43, 94) if "Chapter" in item else RGBColor(38, 38, 38)
        
    doc.add_page_break()
    
    # ------------------ 3. PARSE OUTLINE DOCX PARAGRAPHS ------------------
    doc_outline_path = r"c:\NMPA final project\NMPA_Project_Report_Expanded.docx"
    doc_outline = docx.Document(doc_outline_path)
    
    # Read outline text
    skip_outline = True
    current_chapter = ""
    
    for p in doc_outline.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
            
        if "Chapter 1:" in txt:
            skip_outline = False
            
        if skip_outline:
            continue
            
        # Parse headings vs body text
        if txt.startswith("Chapter "):
            current_chapter = txt
            doc.add_page_break()
            parts = txt.split(":", 1)
            add_styled_heading(doc, parts[0].strip() + ": " + (parts[1].strip() if len(parts) > 1 else ""), 1)
        elif txt.startswith("1.") or txt.startswith("2.") or txt.startswith("3.") or txt.startswith("4.") or txt.startswith("5.") or txt.startswith("6.") or txt.startswith("7."):
            parts = txt.split(" ", 1)
            num = parts[0].strip()
            title = parts[1].strip() if len(parts) > 1 else ""
            dot_count = num.count(".")
            if dot_count == 1:
                add_styled_heading(doc, num + " " + title, 2)
            elif dot_count == 2:
                add_styled_heading(doc, num + " " + title, 3)
            else:
                add_styled_paragraph(doc, txt)
        elif txt.startswith("Server Side:") or txt.startswith("Client Side:") or txt.startswith("Server OS:") or txt.startswith("Database:") or txt.startswith("Backend:") or txt.startswith("Client Browser:"):
            parts = txt.split(":", 1)
            add_styled_bullet(doc, parts[0] + ":", parts[1])
        elif txt.startswith("-") or txt.startswith("*"):
            add_styled_bullet(doc, "", txt[1:].strip())
        else:
            add_styled_paragraph(doc, txt)
            
        # Inject UI Gallery under Chapter 5 at the end
        if "5.3 Bilingual Certification & Offline LocalSync adapters" in txt:
            add_styled_heading(doc, "5.4 User Interface Gallery & Reference Screenshots", 2)
            add_styled_paragraph(doc, "This section presents the high-resolution user interface captures illustrating the visual state changes, direct escalations, bilingual operations, and CAPTCHA workflows implemented in the final production release of the NMPA-CIS platform.")
            
            # Form Portal Screenshot
            add_styled_heading(doc, "5.4.1 Grievance Portal - Form Submission Panel", 3)
            add_styled_paragraph(doc, "Figure 5.1 illustrates the main grievance submission interface. The form provides direct input validation, automatic dropdown bindings for states, bilingual translations switcher, and visual CAPTCHA rendering for public verification security compliance.")
            form_img = r"C:\Users\Inchara salian\.gemini\antigravity-ide\brain\c1e8b3ca-7d2e-4931-bafb-14c4751adcca\grievance_portal_initial_1784605603147.png"
            add_ui_screenshot(doc, form_img, "Figure 5.1: Grievance Registration Interface featuring CAPTCHA and Bilingual Switches")
            
            # Success Modal Screenshot
            add_styled_heading(doc, "5.4.2 Submission Success Validation Modal", 3)
            add_styled_paragraph(doc, "Upon sending a validated grievance request, the interface intercepts the callback response and triggers a native success Modal prompt presenting the unique registration reference string (e.g. NMPA-GRV-XXXX) for subsequent search queries as seen in Figure 5.2.")
            success_img = r"C:\Users\Inchara salian\.gemini\antigravity-ide\brain\c1e8b3ca-7d2e-4931-bafb-14c4751adcca\after_submit_message_1784605954560.png"
            add_ui_screenshot(doc, success_img, "Figure 5.2: Ant Design Grievance Submission Success Notification Modal")
            
            # Standard Queue Screenshot
            add_styled_heading(doc, "5.4.3 Standard Grievance Queue & SLA Timers Console", 3)
            add_styled_paragraph(doc, "Figure 5.3 shows the Platform Administrator's operator console. Complainant details are parsed and organized into standardized compact boxes (Role, Name, Contact, State, Pin) to optimize table space, and real-time SLA countdown countdowns alert personnel of pending timers.")
            queue_img = r"C:\Users\Inchara salian\.gemini\antigravity-ide\brain\c1e8b3ca-7d2e-4931-bafb-14c4751adcca\std_grievance_queue_scrolled_1784608386111.png"
            add_ui_screenshot(doc, queue_img, "Figure 5.3: Operator Dashboard - Standard Grievance Queue with Compact Tags & SLA countdowns")
            
            # Chairman Inbox Screenshot
            add_styled_heading(doc, "5.4.4 Secure Chairman's Office Inbox (Direct Escalations)", 3)
            add_styled_paragraph(doc, "Figure 5.4 displays the secure direct escalation portal accessible strictly by authorized board office terminals. Any grievance flagged for direct escalation, or falling into an SLA breach condition, is routed directly here for immediate administrative action.")
            chair_img = r"C:\Users\Inchara salian\.gemini\antigravity-ide\brain\c1e8b3ca-7d2e-4931-bafb-14c4751adcca\final_chairman_inbox_verification_1784608606545.png"
            add_ui_screenshot(doc, chair_img, "Figure 5.4: Board Office Dashboard - Secure Chairman's Inbox with SLA Breach Routing")
            
    # ------------------ 4. CODE APPENDIX (CHAPTER 8) ------------------
    doc.add_page_break()
    add_styled_heading(doc, "Chapter 8: Source Code Appendix (Comprehensive Code Archive)", 1)
    add_styled_paragraph(doc, "This appendix contains the complete, production-ready source code files implementing the security verification rules, local storage synchronization layers, bilingual translation tables, and Flask backend endpoints. File paths correspond to the active workspace on the local platform.")
    
    code_files = [
        ("backend/app.py", r"c:\NMPA final project\backend\app.py"),
        ("frontend/src/localDb.js", r"c:\NMPA final project\frontend\src\localDb.js"),
        ("frontend/src/pages/GrievancePortal.jsx", r"c:\NMPA final project\frontend\src\pages\GrievancePortal.jsx"),
        ("frontend/src/pages/SystemAdmin.jsx", r"c:\NMPA final project\frontend\src\pages\SystemAdmin.jsx"),
        ("frontend/src/pages/Dashboard.jsx", r"c:\NMPA final project\frontend\src\pages\Dashboard.jsx"),
        ("frontend/src/pages/PortAuthority.jsx", r"c:\NMPA final project\frontend\src\pages\PortAuthority.jsx")
    ]
    
    for label, filepath in code_files:
        if not os.path.exists(filepath):
            continue
            
        add_styled_heading(doc, "8.X File Source: " + label, 2)
        
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            
        add_code_block(doc, label, lines)
        
    # ------------------ 5. SIGN-OFF PAGE (CHAPTER 9) ------------------
    doc.add_page_break()
    add_styled_heading(doc, "Chapter 9: System Sign-Off & Compliance Verification", 1)
    add_styled_paragraph(doc, "This chapter serves as the administrative log certifying the complete integration, inspection testing, and final operations deployment of the New Mangalore Port Authority Cargo Inspection System (NMPA-CIS) v2.0.")
    
    add_styled_heading(doc, "9.1 Operational Integration Checklist", 2)
    add_styled_paragraph(doc, "[X] Port Authority Multi-Role RBAC desks operational and verified.")
    add_styled_paragraph(doc, "[X] Automatic 72-hour Service Level Agreement (SLA) countdown validation verified.")
    add_styled_paragraph(doc, "[X] Secure direct escalation pathway to Chairman's Office Office verified.")
    add_styled_paragraph(doc, "[X] Central billing links and public QR verification verified.")
    add_styled_paragraph(doc, "[X] Local storage synchronisation failover checks verified under offline simulations.")
    add_styled_paragraph(doc, "[X] Database WAL transactional logging and query timeouts certified.")
    
    doc.add_paragraph() # spacing
    add_styled_heading(doc, "9.2 Execution Approval signatures", 2)
    
    # Signature Table (3 cols)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Prepared By:"
    hdr_cells[1].text = "Reviewed By:"
    hdr_cells[2].text = "Approved By:"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    row_cells1 = table.rows[1].cells
    row_cells1[0].text = "\n\n\n______________________"
    row_cells1[1].text = "\n\n\n______________________"
    row_cells1[2].text = "\n\n\n______________________"
    
    row_cells2 = table.rows[2].cells
    row_cells2[0].text = "System Development Lead\nIT Development Cell\nNMPA Port Office"
    row_cells2[1].text = "Quality Assurance Lead\nSecurity Audit Division\nNMPA Port Office"
    row_cells2[2].text = "Port Trust Chairman\nBoard of Trustees\nNew Mangalore Port Authority"
    
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                if p.runs:
                    p.runs[0].font.name = 'Times New Roman'
                    p.runs[0].font.size = Pt(9.5)
                    
    # Save document
    output_path = os.path.abspath("NMPA_Project_Report_Final.docx")
    doc.save(output_path)
    print(f"Successfully generated DOCX: {output_path}")

if __name__ == '__main__':
    generate_report_docx()
